"""
Export Service for Document Generation
Handles PDF and DOCX export with professional formatting
"""
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from pathlib import Path
from datetime import datetime
from app.config import settings
import re


class ExportService:
    """Service for exporting documents to PDF/DOCX with professional formatting"""
    
    def __init__(self):
        """Initialize export service"""
        self.output_dir = Path(settings.GENERATED_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _add_page_number(self, section):
        """Add page numbers to footer"""
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        return footer_para
    
    def _parse_table(self, table_text: str):
        """Parse table from markdown-like format"""
        lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
        
        # Extract table title if present
        title = None
        if lines[0].startswith('TABLE:'):
            title = lines[0].replace('TABLE:', '').strip()
            lines = lines[1:]
        
        # Parse table rows
        rows = []
        for line in lines:
            if line.startswith('|') and line.endswith('|'):
                # Remove leading/trailing pipes and split
                cells = [cell.strip() for cell in line[1:-1].split('|')]
                # Skip separator rows (containing only dashes)
                if not all(set(cell.replace('-', '').strip()) == set() for cell in cells):
                    rows.append(cells)
        
        return title, rows
    
    def export_to_pdf(self, content: str, title: str, filename: str, document_type: str = "document") -> str:
        """
        Export content to PDF with headers, footers, and page numbers
        
        Args:
            content: Document content (markdown-like text)
            title: Document title
            filename: Output filename
            document_type: Type of document
            
        Returns:
            Path to generated PDF file
        """
        output_path = self.output_dir / filename
        
        # Create PDF with narrower margins
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=54,  # 0.75 inch
            leftMargin=54,   # 0.75 inch
            topMargin=90,
            bottomMargin=72
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.black,
            spaceAfter=12,
            alignment=1,  # Center
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.black,
            spaceAfter=6,
            spaceBefore=8,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#E8E8E8'),  # Light grey highlight
            leftIndent=0,
            borderPadding=4
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=4,
            spaceBefore=6,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#F0F0F0'),  # Lighter grey highlight
            leftIndent=20,  # Indented
            borderPadding=4
        )
        
        subsubheading_style = ParagraphStyle(
            'CustomSubSubHeading',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.black,
            spaceAfter=4,
            spaceBefore=6,
            fontName='Helvetica-Bold',
            leftIndent=40  # More indented
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            alignment=4,  # Justified
            spaceAfter=6,
            leading=14,
            fontName='Helvetica',
            leftIndent=0
        )
        
        bullet_style = ParagraphStyle(
            'BulletStyle',
            parent=body_style,
            leftIndent=20,
            bulletIndent=10
        )
        
        sub_bullet_style = ParagraphStyle(
            'SubBulletStyle',
            parent=body_style,
            leftIndent=40,
            bulletIndent=30
        )
        
        # Add title
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 0.15 * inch))
        
        # Add date
        date_text = f"Generated on: {datetime.now().strftime('%B %d, %Y')}"
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=1,
            fontName='Helvetica'
        )
        elements.append(Paragraph(date_text, date_style))
        elements.append(Spacer(1, 0.2 * inch))
        
        # Track current heading level for indentation
        current_level = 0
        
        # Parse and add content
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                elements.append(Spacer(1, 0.05 * inch))
                i += 1
                continue
            
            # Check for table
            if line.startswith('TABLE:') or (line.startswith('|') and '|' in line[1:]):
                # Collect table lines
                table_lines = []
                while i < len(lines) and (lines[i].strip().startswith('|') or lines[i].strip().startswith('TABLE:')):
                    table_lines.append(lines[i])
                    i += 1
                
                # Parse and add table
                table_title, table_data = self._parse_table('\n'.join(table_lines))
                
                if table_title:
                    table_title_style = ParagraphStyle(
                        'TableTitle',
                        parent=styles['Normal'],
                        fontSize=11,
                        fontName='Helvetica-Bold',
                        spaceAfter=4
                    )
                    elements.append(Paragraph(table_title, table_title_style))
                    elements.append(Spacer(1, 0.05 * inch))
                
                if table_data:
                    # Create table with simple outline style
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        # Simple black outline
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        # Header row - bold text, no background color
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('TOPPADDING', (0, 0), (-1, 0), 8),
                        # Data rows
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('TOPPADDING', (0, 1), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    elements.append(t)
                    elements.append(Spacer(1, 0.1 * inch))
                continue
            
            # Clean the line - remove markdown artifacts
            cleaned_line = self._clean_markdown(line)
            
            # Check if it's a heading
            if line.startswith('# '):
                current_level = 1
                elements.append(Paragraph(cleaned_line[2:], heading_style))
            elif line.startswith('## '):
                current_level = 2
                elements.append(Paragraph(cleaned_line[3:], subheading_style))
            elif line.startswith('### '):
                current_level = 3
                elements.append(Paragraph(cleaned_line[4:], subsubheading_style))
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point with indentation based on level
                bullet_text = cleaned_line[2:] if cleaned_line.startswith(('- ', '* ')) else cleaned_line
                if current_level >= 2:
                    elements.append(Paragraph(f"• {bullet_text}", sub_bullet_style))
                else:
                    elements.append(Paragraph(f"• {bullet_text}", bullet_style))
            elif re.match(r'^\d+\.\s+\d+\.\s+', line):
                # Handle incorrect numbering like "10. 1. Text" - extract just the text
                text = re.sub(r'^\d+\.\s+\d+\.\s+', '', cleaned_line)
                if current_level >= 2:
                    elements.append(Paragraph(f"• {text}", sub_bullet_style))
                else:
                    elements.append(Paragraph(f"• {text}", bullet_style))
            elif re.match(r'^\d+\.\s+', line):
                # Numbered list - remove the number, use bullet with indentation
                text = re.sub(r'^\d+\.\s+', '', cleaned_line)
                if current_level >= 2:
                    elements.append(Paragraph(f"• {text}", sub_bullet_style))
                else:
                    elements.append(Paragraph(f"• {text}", bullet_style))
            else:
                # Regular paragraph - justified with indentation based on level
                if current_level >= 2:
                    indented_body = ParagraphStyle(
                        'IndentedBody',
                        parent=body_style,
                        leftIndent=20
                    )
                    elements.append(Paragraph(cleaned_line, indented_body))
                else:
                    elements.append(Paragraph(cleaned_line, body_style))
            
            i += 1
        
        # Build PDF
        doc.build(elements)
        
        return str(output_path)
    
    def _clean_markdown(self, text: str) -> str:
        """Clean markdown artifacts from text"""
        # Remove ** (bold markers)
        text = re.sub(r'\*\*', '', text)
        # Remove * (italic markers) but not bullet points at start
        if not text.startswith('* '):
            text = re.sub(r'\*', '', text)
        # Remove extra spaces
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def export_to_docx(self, content: str, title: str, filename: str, document_type: str = "document") -> str:
        """
        Export content to DOCX with headers, footers, page numbers, and justified text
        
        Args:
            content: Document content (markdown-like text)
            title: Document title
            filename: Output filename
            document_type: Type of document
            
        Returns:
            Path to generated DOCX file
        """
        output_path = self.output_dir / filename
        
        # Create document
        doc = DocxDocument()
        
        # Set up sections with narrower margins
        section = doc.sections[0]
        section.left_margin = Inches(0.75)   # Reduced from 1 inch
        section.right_margin = Inches(0.75)  # Reduced from 1 inch
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        
        # Add header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = title
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.runs[0]
        header_run.font.size = Pt(10)
        header_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        header_run.font.bold = True
        
        # Add a line below header
        header_para.paragraph_format.border_bottom = True
        
        # Add footer with page numbers
        footer_para = self._add_page_number(section)
        footer_para.runs[0].font.size = Pt(10)
        footer_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Add document title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        title_run.font.size = Pt(16)
        
        # Reduce spacing after title
        title_para.paragraph_format.space_after = Pt(6)
        
        # Add date
        date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        date_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Reduce spacing after date
        date_para.paragraph_format.space_after = Pt(12)
        
        # Parse and add content
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Add minimal spacing for empty lines
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(3)
                i += 1
                continue
            
            # Check for table
            if line.startswith('TABLE:') or (line.startswith('|') and '|' in line[1:]):
                # Collect table lines
                table_lines = []
                while i < len(lines) and (lines[i].strip().startswith('|') or lines[i].strip().startswith('TABLE:')):
                    table_lines.append(lines[i])
                    i += 1
                
                # Parse and add table
                table_title, table_data = self._parse_table('\n'.join(table_lines))
                
                if table_title:
                    table_title_para = doc.add_paragraph(self._clean_markdown(table_title))
                    table_title_para.runs[0].font.bold = True
                    table_title_para.runs[0].font.size = Pt(11)
                    table_title_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black
                    table_title_para.paragraph_format.space_after = Pt(4)
                
                if table_data and len(table_data) > 0:
                    # Ensure all rows have the same number of columns
                    max_cols = max(len(row) for row in table_data)
                    normalized_data = []
                    for row in table_data:
                        # Pad rows with empty strings if they have fewer columns
                        normalized_row = row + [''] * (max_cols - len(row))
                        normalized_data.append(normalized_row)
                    
                    # Create table with simple outline style
                    table = doc.add_table(rows=len(normalized_data), cols=max_cols)
                    table.style = 'Table Grid'  # Simple grid style
                    
                    # Fill table
                    for row_idx, row_data in enumerate(normalized_data):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = self._clean_markdown(cell_data)
                                
                                # Format header row - bold, no background color
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.size = Pt(11)
                                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                else:
                                    # Data rows
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)
                                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add minimal space after table
                    space_para = doc.add_paragraph()
                    space_para.paragraph_format.space_after = Pt(6)
                continue
            
            # Clean the line - remove markdown artifacts
            cleaned_line = self._clean_markdown(line)
            
            # Check if it's a heading
            if line.startswith('# '):
                heading = doc.add_heading(cleaned_line[2:], level=1)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black
                heading.runs[0].font.size = Pt(14)
                heading.paragraph_format.space_before = Pt(8)
                heading.paragraph_format.space_after = Pt(6)
            elif line.startswith('## '):
                heading = doc.add_heading(cleaned_line[3:], level=2)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black
                heading.runs[0].font.size = Pt(12)
                heading.paragraph_format.space_before = Pt(6)
                heading.paragraph_format.space_after = Pt(4)
            elif line.startswith('### '):
                heading = doc.add_heading(cleaned_line[4:], level=3)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black
                heading.runs[0].font.size = Pt(11)
                heading.paragraph_format.space_before = Pt(6)
                heading.paragraph_format.space_after = Pt(4)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                bullet_text = cleaned_line[2:] if cleaned_line.startswith(('- ', '* ')) else cleaned_line
                para = doc.add_paragraph(bullet_text, style='List Bullet')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(3)
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            elif re.match(r'^\d+\.\s+\d+\.\s+', line):
                # Handle incorrect numbering like "10. 1. Text" - extract just the text
                text = re.sub(r'^\d+\.\s+\d+\.\s+', '', cleaned_line)
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(3)
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            elif re.match(r'^\d+\.\s+', line):
                # Numbered list - remove the number, use bullet instead
                text = re.sub(r'^\d+\.\s+', '', cleaned_line)
                para = doc.add_paragraph(text, style='List Bullet')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(3)
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            else:
                # Regular paragraph - JUSTIFIED
                para = doc.add_paragraph(cleaned_line)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                para.paragraph_format.space_after = Pt(6)  # Reduced spacing
                
                # Set font
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            
            i += 1
        
        # Save document
        doc.save(str(output_path))
        
        return str(output_path)
    
    def export_document(
        self,
        content: str,
        title: str,
        document_id: str,
        format: str = "pdf",
        document_type: str = "document"
    ) -> str:
        """
        Export document in specified format
        
        Args:
            content: Document content
            title: Document title
            document_id: Unique document ID
            format: Output format (pdf or docx)
            document_type: Type of document
            
        Returns:
            Path to generated file
        """
        # Generate filename - sanitize title properly
        # Remove newlines, tabs, and other whitespace characters first
        clean_title = title.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
        # Remove invalid filename characters
        safe_title = re.sub(r'[^\w\s-]', '', clean_title).strip()
        # Replace multiple spaces with single space
        safe_title = re.sub(r'\s+', ' ', safe_title)
        # Replace spaces with underscores
        safe_title = safe_title.replace(' ', '_')
        # Limit length to avoid path too long errors
        safe_title = safe_title[:100]
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{safe_title}_{timestamp}.{format}"
        
        if format.lower() == 'pdf':
            return self.export_to_pdf(content, title, filename, document_type)
        elif format.lower() in ['docx', 'doc']:
            return self.export_to_docx(content, title, filename, document_type)
        else:
            raise ValueError(f"Unsupported format: {format}")


# Global instance
export_service = ExportService()
